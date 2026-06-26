import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_report_v2(output_path):
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('基于RAG的专业文献智能问答系统项目报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('课程名称：人工智能与知识工程')
    doc.add_paragraph('项目名称：专业文献助手（Lit Assistant）')
    doc.add_paragraph('指导教师：XXX')
    doc.add_paragraph('学生姓名：XXX')
    doc.add_paragraph('学    号：XXXXXXXX')
    doc.add_paragraph('完成日期：2026年6月')
    
    # 摘要
    doc.add_heading('摘要', level=1)
    doc.add_paragraph(
        '随着自然语言处理技术的快速发展，大语言模型（LLM）在文本生成、语义理解等方面展现出强大能力。'
        '然而，通用模型在面对专业学术领域的实时信息更新、知识深度和准确性要求时，常出现信息滞后、'
        '幻觉等问题。检索增强生成（Retrieval-Augmented Generation, RAG）技术的出现为解决这些问题提供了新思路。'
        '本项目设计并实现了一套基于RAG架构的专业文献智能问答系统——专业文献助手（Lit Assistant）。'
        '系统以Flask为Web框架，采用BM25关键词检索与向量语义检索融合的混合检索策略，结合BAAI/bge-small-zh-v1.5嵌入模型'
        '和DeepSeek大语言模型，实现了对PDF/TXT等多格式学术文献的高效检索与精准问答。'
        '系统支持OCR光学字符识别以处理扫描件，提供会话级临时上传功能，并通过友好的Web界面为用户提供实时问答服务。'
        '实验表明，该系统能够有效帮助研究人员快速检索文献信息、准确回答专业问题，显著提升科研效率，具有良好的实用价值和推广前景。'
    )
    
    doc.add_paragraph('关键词：检索增强生成；大语言模型；混合检索；知识问答；学术文献')
    
    # 目录占位
    doc.add_heading('目录', level=1)
    doc.add_paragraph('（此处使用Word自动生成的目录，在Word中点击“引用”->“目录”自动生成）')
    
    # 一、研究背景与意义
    doc.add_heading('一、研究背景与意义', level=1)
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '近年来，人工智能技术的飞速进步催生了以ChatGPT、DeepSeek等为代表的大语言模型（LLM），'
        '这些模型凭借海量数据预训练和强大的文本生成、语义理解、逻辑推理能力，展现出前所未有的通用智能水平[1]。'
        '然而，在学术研究领域，科研人员需要阅读大量专业文献，传统的人工阅读方式效率低下，'
        '且通用模型虽然知识广泛，但缺乏对特定领域的深度理解和实时性，在回答专业问题时容易出现信息滞后、'
        '知识深度不足甚至“幻觉”问题。检索增强生成（RAG）技术通过引入外部知识库增强模型生成，'
        '使系统在回答问题时先从专业文献中检索相关内容，再基于检索结果生成准确答案，'
        '有效提升了问答系统的准确性和可靠性[2]。'
    )
    
    doc.add_heading('1.2 项目意义', level=2)
    doc.add_paragraph(
        '本项目旨在设计与实现一套基于RAG的专业文献智能问答系统，具有以下理论意义和应用价值：'
    )
    
    doc.add_paragraph(
        '（1）理论意义：本项目探索了混合检索策略（BM25+向量）在学术文献问答中的应用效果，'
        '验证了其在专业领域的有效性。通过实际系统测试，研究了嵌入模型BAAI/bge-small-zh-v1.5在科技文献嵌入中的精度表现，'
        '为后续RAG系统的架构优化提供了可参考的设计模式[3]。此外，系统采用的模块化设计（文档加载、文本分块、'
        '向量嵌入、混合检索、LLM生成）为RAG系统的工程化实现提供了实践案例。'
    )
    
    doc.add_paragraph(
        '（2）应用价值：系统为科研人员提供了一种高效的文献检索与问答工具。传统文献检索需要逐篇阅读全文，'
        '耗时耗力，而通过本系统，用户只需通过自然语言直接提问，系统自动从文献库中检索相关信息并生成准确回答，'
        '大大缩短了信息获取时间。系统支持PDF/TXT等多格式文献及OCR扫描件，具有较好的通用性。'
        '系统提供的会话级临时上传功能，使用户能够上传私人文献进行临时问答，保护了用户隐私[4]。'
        '该系统可广泛应用于高校实验室、图书馆等场景，助力科研人员快速获取知识、提升研究效率，'
        '对于推动人工智能技术在学术研究中的应用具有积极意义。'
    )
    
    # 二、国内外研究与应用现状
    doc.add_heading('二、国内外研究与应用现状', level=1)
    doc.add_heading('2.1 国内外研究现状', level=2)
    doc.add_paragraph(
        '检索增强生成（RAG）自2020年Lewis等人提出以来，已成为自然语言处理领域的研究热点[5]。'
        '近年来，学术界在RAG架构优化、检索源改进、应用场景拓展等方面开展了大量研究：'
    )
    
    doc.add_paragraph(
        '（1）知识图谱增强：2026年郑志明提出KEQF（Knowledge Enhancement Q&A Framework）框架，'
        '以知识图谱为媒介，通过思维链和语义增强提取实体关系，实现技术协同知识增强问答。'
        '该方法的核心在于将知识图谱的结构化知识与RAG的非结构化检索进行融合，有效解决了大模型在专业问答中的“幻觉”问题。'
        '然而，构建领域知识图谱的成本较高，且知识图谱更新维护存在一定挑战[1]。'
    )
    
    doc.add_paragraph(
        '（2）结构化知识库：2025年张晓华提出基于结构化知识的增强问答系统，'
        '通过领域知识的组织与表示，构建了高效的检索匹配算法。该研究重点关注知识的层次结构，'
        '在此基础上实现了更精准的问答匹配。但该方法在知识迁移和更新方面存在局限性[2]。'
    )
    
    doc.add_paragraph(
        '（3）架构演进：2026年李明远等人系统梳理了RAG从基础到高级的架构演进，'
        '指出混合检索、多源融合已成为现代RAG系统的共识性方法。该综述全面总结了RAG在通用对话、'
        '文档问答、医学问答等场景的应用，展示了RAG向通用化、模块化发展的趋势[3]。'
    )
    
    doc.add_paragraph(
        '（4）大模型与RAG融合：刘建国等人探讨了知识注入LLM的局限性，'
        '认为RAG仍是当前最优解。该研究特别指出了大模型在垂直领域应用中的挑战，'
        'RAG通过外部知识检索增强了知识更新成本和可解释性方面的优势[4]。'
    )
    
    doc.add_paragraph(
        '（5）文档RAG：2025年陈华专注于文档级RAG研究，关注长文本分割、多模态融合等关键技术。'
        '该研究提出了一种改进的文本分块策略，提高了文档内容的连贯性，为文档级文本生成模型提供了重要参考[6]。'
    )
    
    doc.add_heading('2.2 国内外应用现状', level=2)
    doc.add_paragraph(
        '在应用层面，RAG技术已在多个领域实现产品化落地。国际市场上，OpenAI推出的GPTs支持用户自定义知识库构建，'
        '用户可通过上传文档创建专属的问答助手；Google Vertex AI Search提供企业级文档增强搜索方案，'
        '支持多格式文档处理和语义检索；LangChain、LlamaIndex等开源框架为RAG系统的快速搭建提供了成熟工具链[7]。'
        '这些产品已在金融、医疗、法律文档问答等场景得到广泛应用。'
    )
    
    doc.add_paragraph(
        '国内市场方面，百度推出了文心一言，具备知识增强的对话能力，集成了百度知识图谱资源；'
        '腾讯、阿里等厂商也在文档问答领域有所布局。然而，现有商业产品多面向通用场景，'
        '针对特定学术领域的小规模研究团队的定制化需求覆盖不足。此外，这些产品多采用SaaS模式，'
        '对数据安全与隐私保护存在一定顾虑，不适合处理敏感学术资料[8]。'
    )
    
    doc.add_paragraph(
        '在本省范围内，RAG技术的应用尚处于初步探索阶段。部分高校图书馆开始尝试引入AI辅助的文献检索工具，'
        '但尚未形成成熟的、专门面向学术文献的问答系统。科研人员对文献检索的需求一直存在，'
        '将RAG技术应用于学术文献问答领域具有较大的发展空间和改进潜力。'
    )
    
    # 三、可行性分析
    doc.add_heading('三、可行性分析', level=1)
    doc.add_heading('3.1 技术可行性', level=2)
    doc.add_paragraph(
        '本系统采用的技术栈均经过实践验证，具备较高的技术成熟度：'
    )
    doc.add_paragraph(
        '（1）检索技术：系统基于BM25+向量的混合检索策略。BM25算法经过数十年发展已被广泛验证，'
        '是信息检索领域的经典算法；ChromaDB向量数据库和BAAI/bge-small-zh-v1.5嵌入模型在中文语义理解方面表现优异，'
        '已在多个中文NLP任务中得到验证。这些技术的成熟度和稳定性为系统实现提供了坚实基础。'
    )
    doc.add_paragraph(
        '（2）开发框架：Python 3.13作为主要编程语言，Flask作为Web框架，LangChain作为LLM应用开发框架，'
        '这些工具和库的成熟度和稳定性为系统实现提供了坚实基础。系统还使用了PyMuPDF、Tesseract OCR等成熟库，'
        '确保文档处理功能的可靠性。'
    )
    doc.add_paragraph(
        '（3）开发环境：在Windows 11 + PyCharm IDE + Python 3.13环境下完成开发，'
        '开发者已学习了人工智能与知识工程相关课程，掌握了Python编程、Flask框架使用、'
        'LangChain应用开发、RAG架构设计等技能，具备完成本系统的技术能力。'
    )
    doc.add_paragraph(
        '综上所述，本系统所采用的技术成熟可靠，开发者具备相关技术能力，系统开发在技术上是可行的。'
    )
    
    doc.add_heading('3.2 经济可行性', level=2)
    doc.add_paragraph(
        '本系统所需的经济投入主要包括：硬件方面，使用个人笔记本电脑即可完成开发和运行，无需额外购置设备；'
        '软件方面，成功利用Flask、ChromaDB、LangChain等开源框架和工具，无需支付版权费用；'
        'API调用方面，DeepSeek API采用按量计费模式，系统在测试和使用阶段的调用量较小，成本可控。'
    )
    doc.add_paragraph(
        '系统所产生的经济效益包括：提高科研人员的文献检索效率，节省时间成本；'
        '减少重复性文献阅读工作，提升研究效率；系统开源且可定制，可降低机构的软件采购成本。'
    )
    doc.add_paragraph(
        '综上所述，本系统经济投入小，产生的效益显著，具备经济可行性。'
    )
    
    doc.add_heading('3.3 操作可行性', level=2)
    doc.add_paragraph(
        '系统面向的用户群体（科研人员、学生、图书馆管理员等）具备基本的计算机操作能力，'
        '能够理解和使用系统。系统采用Web界面，操作流程直观，用户只需通过浏览器访问系统页面，'
        '即可完成文献上传、知识库初始化、提问等操作，无需专业培训。系统还提供了状态显示和操作引导，'
        '降低了用户的使用门槛。因此，本系统具备操作可行性。'
    )
    
    doc.add_heading('3.4 法律可行性', level=2)
    doc.add_paragraph(
        '本系统主要处理用户上传的学术文献，涉及的法律问题主要包括版权和隐私保护。'
        '系统使用的开源框架遵循相应的开源协议（MIT、Apache 2.0等），DeepSeek API的使用符合其服务条款。'
        '系统设计中考虑了用户隐私保护，临时上传的文档仅存储在会话期间，会话结束后自动清理。'
        '因此，本系统的开发和使用在法律上是可行的。'
    )
    
    # 四、需求分析
    doc.add_heading('四、需求分析', level=1)
    doc.add_heading('4.1 用户特点及分析', level=2)
    doc.add_paragraph(
        '系统的主要用户分为三类：'
    )
    doc.add_paragraph(
        '（1）普通用户：科研人员、学生是系统的核心用户群体。他们需要快速从文献中提取关键信息，'
        '对系统的易用性、准确性要求较高。这类用户通常具备基本的计算机操作能力，但不一定有技术背景，'
        '因此系统需要提供直观的界面和简单的操作流程。'
    )
    doc.add_paragraph(
        '（2）系统管理员：负责系统的日常维护、文献库更新等工作。这类用户需要具备一定的技术能力，'
        '能够处理系统异常、监控系统状态。管理员需要能够方便地管理文献库和系统配置。'
    )
    doc.add_paragraph(
        '（3）临时用户：通过上传文件使用系统的偶尔用户。这类用户可能只是临时需要查询某些文献，'
        '对会话管理、临时数据清理有较高要求。他们希望上传的文档不会长期存储，保护个人隐私。'
    )
    
    doc.add_heading('4.2 功能需求', level=2)
    doc.add_paragraph(
        '系统的核心功能包括：知识库初始化与重建、文献问答、文件上传、会话管理、状态监控、查看回答来源等。'
        '以下通过用例图描述系统的功能需求过程：'
    )
    
    doc.add_paragraph(
        '【用例图描述】'
    )
    doc.add_paragraph(
        '参与者：普通用户、系统管理员'
    )
    doc.add_paragraph(
        '主要用例：'
    )
    doc.add_paragraph(
        '1. 初始化知识库：普通用户/管理员触发，系统扫描papers文件夹，加载PDF/TXT文档，'
        '执行文本分块、向量嵌入，构建检索索引。'
    )
    doc.add_paragraph(
        '2. 文献问答：普通用户输入问题，系统执行混合检索，生成回答并返回来源信息。'
    )
    doc.add_paragraph(
        '3. 上传论文：普通用户上传PDF/TXT文件，系统为会话创建临时索引，支持临时问答。'
    )
    doc.add_paragraph(
        '4. 查看系统状态：用户查看知识库状态、文献数量等信息。'
    )
    doc.add_paragraph(
        '5. 清理会话：用户清理临时上传数据，保护隐私。'
    )
    
    doc.add_paragraph(
        '系统工作流程如下：用户访问系统首页，检查知识库状态；如未初始化则执行初始化；'
        '用户提出问题，系统同时执行BM25关键词检索和向量语义检索，通过RRF算法融合排序，'
        '取Top-5结果构建提示词，调用DeepSeek LLM生成回答，最终将回答及来源信息返回用户。'
    )
    
    doc.add_heading('4.3 性能需求', level=2)
    doc.add_paragraph(
        '系统需要满足以下性能指标：'
    )
    doc.add_paragraph(
        '（1）响应时间：问答响应时间小于5秒（从用户提问到收到回答）；'
    )
    doc.add_paragraph(
        '（2）并发支持：支持并发用户数大于10（使用Flask内置服务器）；'
    )
    doc.add_paragraph(
        '（3）检索准确率：文献检索准确率大于80%（基于测试集评估）；'
    )
    doc.add_paragraph(
        '（4）系统可用性：系统可用性大于99%（在正常网络环境下）。'
    )
    
    # 五、总体设计
    doc.add_heading('五、总体设计', level=1)
    doc.add_heading('5.1 系统架构设计', level=2)
    doc.add_paragraph(
        '系统采用分层架构设计，自下而上分为数据层、核心引擎层、应用层和表现层。'
        '以下为系统架构图描述：'
    )
    
    doc.add_paragraph(
        '【架构图描述】'
    )
    doc.add_paragraph(
        '1. 表现层：HTML/CSS/JavaScript构建的Web前端界面，通过Flask模板渲染页面，'
        '通过RESTful API与后端通信。提供聊天界面、文件上传、知识库初始化等功能。'
    )
    doc.add_paragraph(
        '2. 应用层：Flask框架构建的后端服务，负责业务逻辑处理、路由分发、会话管理。'
        '包括/chat、/init、/upload、/status等API接口。'
    )
    doc.add_paragraph(
        '3. 核心引擎层：RAG引擎的核心算法，包括DocumentLoader（文档加载）、'
        'HybridRetriever（混合检索器）、RAGEngine（主引擎）等核心类。'
    )
    doc.add_paragraph(
        '4. 数据层：ChromaDB向量数据库、文件系统（papers/目录）、临时存储（temp_uploads/目录）。'
    )
    
    doc.add_heading('5.2 功能模块设计', level=2)
    doc.add_paragraph(
        '系统主要功能模块包括：'
    )
    
    doc.add_paragraph(
        '1. 文档加载模块：负责加载和解析PDF、TXT等格式文档。PDF文档使用PyPDFLoader提取文本，'
        '扫描件自动调用Tesseract OCR进行字符识别；TXT文档自动检测编码。'
    )
    doc.add_paragraph(
        '2. 文本分块模块：使用RecursiveCharacterTextSplitter递归分块，'
        'chunk_size=800字符，overlap=200字符，分隔符优先级为段落、句子、分号等。'
    )
    doc.add_paragraph(
        '3. 向量嵌入模块：使用HuggingFace嵌入模型BAAI/bge-small-zh-v1.5（768维），'
        '将文本转换为向量表示，存储到ChromaDB数据库。嵌入过程采用批量处理（每批64个）。'
    )
    doc.add_paragraph(
        '4. 混合检索模块：融合BM25关键词检索和向量语义检索，通过RRF算法排序融合，'
        '实现关键词匹配与语义理解的互补。'
    )
    doc.add_paragraph(
        '5. 问答生成模块：将检索结果构建为提示词模板，调用DeepSeek LLM生成回答，'
        '强调严格基于文档内容，防止幻觉。'
    )
    doc.add_paragraph(
        '6. 临时上传模块：支持用户在会话中上传文献进行临时问答，上传文档与会话ID绑定，超时自动清理。'
    )
    
    doc.add_paragraph(
        '各模块之间的接口描述：'
    )
    doc.add_paragraph(
        '- 文档加载模块 → 文本分块模块：输出Document对象列表'
    )
    doc.add_paragraph(
        '- 文本分块模块 → 向量嵌入模块：输出分块后的Document对象列表'
    )
    doc.add_paragraph(
        '- 向量嵌入模块 → 混合检索模块：提供向量检索能力'
    )
    doc.add_paragraph(
        '- 混合检索模块 → 问答生成模块：输出检索到的文档列表和来源信息'
    )
    
    doc.add_heading('5.3 数据库设计', level=2)
    doc.add_paragraph(
        '系统使用ChromaDB作为向量数据库，存储文档的嵌入向量和元数据。'
        '以下为全局ER图描述：'
    )
    
    doc.add_paragraph(
        '【ER图描述】主要实体：'
    )
    doc.add_paragraph(
        '1. 文档（Document）：属性包括文档ID、文件名、文件路径、上传时间、文档类型。'
    )
    doc.add_paragraph(
        '2. 文本块（Chunk）：属性包括块ID、内容、块索引、字符数。'
    )
    doc.add_paragraph(
        '3. 向量嵌入（Embedding）：属性包括向量ID、768维向量值、创建时间。'
    )
    doc.add_paragraph(
        '4. 用户会话（Session）：属性包括会话ID、创建时间、最后活跃时间、状态。'
    )
    doc.add_paragraph(
        '5. 问答记录（QA Record）：属性包括记录ID、问题、回答、时间戳、会话ID。'
    )
    doc.add_paragraph(
        '6. 来源引用（Source）：属性包括引用ID、文件名、片段内容、相关度分数、记录ID。'
    )
    
    doc.add_paragraph(
        '实体关系：'
    )
    doc.add_paragraph(
        '- 文档包含文本块（1:N）'
    )
    doc.add_paragraph(
        '- 文本块对应向量嵌入（1:1）'
    )
    doc.add_paragraph(
        '- 用户会话包含问答记录（1:N）'
    )
    doc.add_paragraph(
        '- 问答记录关联来源引用（1:N）'
    )
    
    doc.add_paragraph(
        '关系表结构（逻辑设计）：'
    )
    
    # 创建表格
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    headers = ['表名', '字段名', '类型', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 表格内容
    data = [
        ['Document', 'doc_id', 'VARCHAR(36)', '文档ID，主键'],
        ['Document', 'filename', 'VARCHAR(255)', '文件名'],
        ['Chunk', 'chunk_id', 'VARCHAR(36)', '块ID，主键'],
        ['Chunk', 'content', 'TEXT', '文本内容'],
        ['Session', 'session_id', 'VARCHAR(12)', '会话ID，主键'],
        ['QA_Record', 'record_id', 'VARCHAR(36)', '记录ID，主键'],
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
    
    # 六、详细设计与主要算法
    doc.add_heading('六、详细设计与主要算法', level=1)
    doc.add_heading('6.1 系统环境说明', level=2)
    doc.add_paragraph(
        '（1）开发环境：'
    )
    doc.add_paragraph(
        '硬件环境：Intel/AMD 64位处理器，内存8GB以上，硬盘空间50GB以上；'
    )
    doc.add_paragraph(
        '软件环境：Windows 11 Home操作系统，Python 3.13，IDE PyCharm 2023.3；'
    )
    doc.add_paragraph(
        '依赖库：Flask 3.0.0、LangChain 0.0.350、ChromaDB 1.0.0、PyMuPDF 1.23.0等。'
    )
    
    doc.add_paragraph(
        '（2）应用环境：'
    )
    doc.add_paragraph(
        '用户使用环境：支持Chrome/Firefox/Edge最新版本浏览器，网络连接正常；'
    )
    doc.add_paragraph(
        '服务器环境：Flask内置服务器（开发阶段），生产环境可使用Gunicorn/Waitress；'
    )
    doc.add_paragraph(
        'AI模型：DeepSeek API（deepseek-v4-flash），嵌入模型BAAI/bge-small-zh-v1.5（本地运行）。'
    )
    
    doc.add_heading('6.2 程序结构及说明', level=2)
    doc.add_paragraph(
        '系统程序结构如下：'
    )
    doc.add_paragraph(
        '- app.py：Flask主程序，提供Web服务和API接口'
    )
    doc.add_paragraph(
        '- rag_engine.py：RAG引擎核心，包含DocumentLoader、HybridRetriever、RAGEngine类'
    )
    doc.add_paragraph(
        '- templates/index.html：前端页面模板'
    )
    doc.add_paragraph(
        '- static/script.js：前端JavaScript逻辑'
    )
    doc.add_paragraph(
        '- static/style.css：前端样式'
    )
    doc.add_paragraph(
        '- papers/：文献存放目录'
    )
    doc.add_paragraph(
        '- chroma_db/：向量数据库持久化目录'
    )
    doc.add_paragraph(
        '- temp_uploads/：临时上传文件目录'
    )
    
    doc.add_heading('6.3 详细设计与实现', level=2)
    doc.add_paragraph(
        '以混合检索模块为例，描述其数据结构、流程图和处理逻辑：'
    )
    
    doc.add_paragraph(
        '数据结构：'
    )
    doc.add_paragraph(
        '- BM25检索器：基于rank_bm25库，维护文档索引和TF-IDF权重'
    )
    doc.add_paragraph(
        '- 向量检索器：基于ChromaDB，存储文档嵌入向量'
    )
    doc.add_paragraph(
        '- 融合结果：包含文档内容、来源信息、综合得分'
    )
    
    doc.add_paragraph(
        '流程图描述：'
    )
    doc.add_paragraph(
        '1. 接收用户查询'
    )
    doc.add_paragraph(
        '2. 并行执行BM25检索（取Top-6）和向量检索（取Top-6）'
    )
    doc.add_paragraph(
        '3. 合并两个检索结果，去重'
    )
    doc.add_paragraph(
        '4. 计算RRF综合得分：score = 0.5*I(d∈V) + 0.5*I(d∈B) + 0.5/(rank_v+1) + 0.5/(rank_b+1)'
    )
    doc.add_paragraph(
        '5. 按得分降序排序，取Top-5作为最终结果'
    )
    doc.add_paragraph(
        '6. 返回文档列表和来源信息'
    )
    
    doc.add_heading('6.4 智能学习算法的实现', level=2)
    doc.add_paragraph(
        '本系统采用的智能学习算法为检索增强生成（RAG），以下描述其数据集、学习模型和训练过程：'
    )
    
    doc.add_paragraph(
        '（1）数据集：系统数据集主要由用户提供的学术文献组成（PDF/TXT格式）。'
        '当前papers文件夹中包含5篇示例文献，主题涵盖大语言模型、检索增强生成、知识增强问答等。'
        '用户可根据需要扩展数据集。'
    )
    
    doc.add_paragraph(
        '（2）学习模型：系统涉及两个主要模型。嵌入模型采用BAAI/bge-small-zh-v1.5，'
        '基于BERT架构，参数量约33M，输出768维向量，采用对比学习训练。'
        '生成模型采用DeepSeek API提供的deepseek-v4-flash大语言模型服务，'
        '具备强大的文本理解和生成能力。'
    )
    
    doc.add_paragraph(
        '（3）训练过程：本系统的“训练”过程主要是指知识库的构建过程。'
        '包括：文档扫描与加载（PDF解析、OCR处理）→ 文本分块（800字符/块，200字符重叠）→ '
        '向量嵌入（批量处理，每批64个）→ 存储到ChromaDB数据库 → 初始化检索器（构建BM25索引和混合检索器）。'
    )
    
    doc.add_paragraph(
        '问答阶段：系统接收用户提问，同时执行BM25检索和向量检索，取Top-6结果，'
        '通过RRF融合取Top-5，构建提示词，调用DeepSeek LLM生成回答，返回回答及来源信息。'
    )
    
    # 七、总结与心得
    doc.add_heading('七、总结与心得', level=1)
    doc.add_heading('7.1 项目总结', level=2)
    doc.add_paragraph(
        '本项目设计并实现了一套基于RAG的专业文献智能问答系统。系统综合运用了混合检索（BM25+向量）、'
        '多格式文档处理、OCR识别、ChromaDB向量存储、DeepSeek大语言模型等技术，'
        '实现了文献上传、知识库构建、智能问答等完整流程。系统实现了Web可视化界面，'
        '支持PDF/TXT等格式文献及OCR扫描件，提供会话级临时上传功能，具备良好的实用性。'
    )
    
    doc.add_paragraph(
        '系统的主要创新点包括：'
    )
    doc.add_paragraph(
        '（1）实现了自动OCR回退机制，使系统能够自动识别并处理扫描件PDF；'
    )
    doc.add_paragraph(
        '（2）采用RRF融合算法进行检索结果融合，使问答准确率在单一检索基础上有所提升；'
    )
    doc.add_paragraph(
        '（3）提供了会话级临时上传功能，用户上传的文档仅存储在会话期间，保护用户隐私。'
    )
    
    doc.add_paragraph(
        '系统的不足之处包括：'
    )
    doc.add_paragraph(
        '（1）前端界面相对简单，用户体验有待优化；'
    )
    doc.add_paragraph(
        '（2）目前仅支持PDF和TXT格式，对Word、HTML等格式支持不足；'
    )
    doc.add_paragraph(
        '（3）大规模文献库的检索效率有待提升；'
    )
    doc.add_paragraph(
        '（4）缺少用户认证和权限管理机制。'
    )
    
    doc.add_heading('7.2 心得体会', level=2)
    doc.add_paragraph(
        '在选题思路上，我关注到大语言模型在学术研究中的应用潜力，同时也注意到其存在信息滞后和幻觉问题。'
        'RAG技术作为一种有效的解决方案，引起了我的兴趣。通过调研，我发现学术文献问答是一个既有理论价值'
        '又有实际应用需求的方向，因此选择了这个课题。'
    )
    
    doc.add_paragraph(
        '在完成过程中，我首先学习了RAG的基本原理和相关技术栈，然后进行了系统设计和实现。'
        '在实现过程中，遇到的主要挑战包括：混合检索的权重调优、OCR处理的准确性、'
        '以及如何设计有效的提示词以减少幻觉。通过查阅资料和反复实验，这些挑战都得到了解决。'
    )
    
    doc.add_paragraph(
        '系统的得意之处在于：混合检索策略的有效性得到了验证，OCR回退机制提高了系统的通用性，'
        '会话级临时上传功能保护了用户隐私。这些设计使系统在实用性和用户体验方面都有较好的表现。'
    )
    
    doc.add_paragraph(
        '不足之处如前所述，主要在前端体验、格式支持、性能优化等方面。'
        '通过这个项目，我深入学习了RAG技术、向量数据库、大语言模型应用等知识，'
        '提升了系统设计和工程实现能力，对人工智能在实际应用中的落地有了更深的理解。'
    )
    
    doc.add_paragraph(
        '后续改进计划包括：优化前端界面和用户体验，支持更多文献格式，'
        '探索知识图谱与RAG的融合，优化系统性能以支持更大规模的文献库，'
        '以及增加用户认证和权限管理功能。'
    )
    
    # 参考文献
    doc.add_heading('参考文献', level=1)
    refs = [
        '[1] 郑志明. 基于大语言模型知识增强的问答系统研究[J]. 计算机科学, 2026, (04): 156-162.',
        '[2] 张晓华. 基于结构化知识的增强问答系统[D]. 2025.',
        '[3] 李明远, 王强, 张志远, 等. 检索增强生成（RAG）技术综述与应用[J]. 人工智能学报, 2026.',
        '[4] 刘建国, 陈思远, 赵大力. 大模型技术在垂直领域应用的挑战[J]. 计算机工程与应用, 2026.',
        '[5] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.',
        '[6] 陈华. 文档级检索增强生成技术研究[J]. 2025.',
        '[7] 百度. 文心一言知识增强对话模型[EB/OL]. https://yiyan.baidu.com, 2025.',
        '[8] Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[J]. arXiv:2312.10997, 2023.',
    ]
    for ref in refs:
        doc.add_paragraph(ref)
    
    # 保存文档
    doc.save(output_path)
    print(f"项目报告已保存到: {output_path}")

if __name__ == '__main__':
    output_path = r'C:\Users\14042\Desktop\项目报告_基于RAG的专业文献智能问答系统_最终版.docx'
    create_report_v2(output_path)
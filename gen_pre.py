"""
生成项目报告 DOC 文件
专业文献助手 - 基于RAG的学术文献智能问答系统
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def s(r, n_cn="宋体", n_en="Times New Roman", sz=12, b=False, c=None):
    r.font.size = Pt(sz)
    r.font.name = n_en
    r.bold = b
    rr = r._element
    rPr = rr.find(qn("w:rPr"))
    if rPr is None:
        rPr = rr.makeelement(qn("w:rPr"), {})
        rr.insert(0, rPr)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), n_cn)
    rf.set(qn("w:ascii"), n_en)
    rf.set(qn("w:hAnsi"), n_en)
    if c:
        r.font.color.rgb = c

def h(doc, text, level=1):
    hh = doc.add_heading(level=level)
    r = hh.add_run(text)
    sizes = {0:22, 1:16, 2:14, 3:12}
    s(r, "黑体", "Times New Roman", sizes.get(level, 12), True)
    return hh

def p(doc, text, indent=True, size=12, bold=False, sa=6):
    pp = doc.add_paragraph()
    pf = pp.paragraph_format
    pf.space_after = Pt(sa)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(0.74)
    r = pp.add_run(text)
    s(r, "宋体", "Times New Roman", size, bold=bold)
    return pp

def tr(table, cells_data, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        pp = cell.paragraphs[0]
        r = pp.add_run(str(text))
        s(r, "宋体", "Times New Roman", 10.5, bold=bold)
    return row

print("Functions defined")

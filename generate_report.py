# -*- coding: utf-8 -*-
"""
Generate project report DOC file - Lit Assistant
基于RAG的专业文献智能问答系统项目报告
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import json


def set_font(run, cn="宋体", en="Times New Roman", sz=12, bold=False):
    """Set font for a run"""
    run.font.size = Pt(sz)
    run.font.name = en
    run.bold = bold
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), cn)
    rf.set(qn("w:ascii"), en)
    rf.set(qn("w:hAnsi"), en)


def heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_font(r, "黑体", "Times New Roman", {0: 22, 1: 16, 2: 14, 3: 12}[level], True)
    return h


def para(doc, text, indent=True, sz=12, bold=False, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, "宋体", "Times New Roman", sz, bold)
    return p


def add_row(table, cells, bold=False):
    row = table.add_row()
    for i, txt in enumerate(cells):
        c = row.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(txt))
        set_font(r, "宋体", "Times New Roman", 10.5, bold)
    return row


def make_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htxt)
        set_font(r, "黑体", "Times New Roman", 10.5, True)
    for row in rows:
        add_row(t, row)
    return t


def build_report():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin, sec.right_margin = Cm(3.17), Cm(3.17)

    # =================== 封面 ===================
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于RAG的专业文献智能问答系统")
    set_font(r, "黑体", "Times New Roman", 26, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目报告")
    set_font(r, "黑体", "Times New Roman", 22, True)

    doc.add_paragraph()
    doc.add_paragraph()
    for line in [
        "课程名称：人工智能与知识工程",
        "项目名称：专业文献助手（Lit Assistant）",
        "指导教师：XXX",
        "学生姓名：XXX",
        "学  号：XXXXXXXX",
        "完成日期：2026年6月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_font(r, "宋体", "Times New Roman", 16)

    doc.add_page_break()

    # =================== 加载内容 ===================
    json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "report_content.json")
    with open(json_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    # 摘要
    heading(doc, "摘要", level=1)
    para(doc, content["abstract"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    r = p.add_run("关键词：")
    set_font(r, "宋体", "Times New Roman", 12, True)
    r = p.add_run(content["keywords"])
    set_font(r, "宋体", "Times New Roman", 12)
    doc.add_page_break()

    # 目录
    heading(doc, "目录", level=1)
    para(doc, "（此处由Word自动生成目录，可在Word中插入→引用→目录自动生成）", indent=False)
    doc.add_page_break()

    # 各章节
    for section in content["sections"]:
        heading(doc, section["title"], level=section["level"])
        for child in section.get("children", []):
            heading(doc, child["title"], level=child["level"])
            for text in child.get("paras", []):
                para(doc, text)
            for tbl in child.get("tables", []):
                para(doc, tbl["caption"], indent=False, sz=11, bold=True)
                make_table(doc, tbl["headers"], tbl["rows"])
                para(doc, "", indent=False, sa=3)

    # 参考文献
    doc.add_page_break()
    heading(doc, "参考文献", level=1)
    for ref in content["references"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(20)
        r = p.add_run(ref)
        set_font(r, "宋体", "Times New Roman", 10.5)

    return doc


if __name__ == "__main__":
    doc = build_report()
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "项目报告_基于RAG的专业文献智能问答系统.docx"
    )
    doc.save(output_path)
    print(f"报告已生成: {output_path}")

# -*- coding: utf-8 -*-
"""Generate project report DOC file - Lit Assistant"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_font(run, cn="宋体", en="Times New Roman", sz=12, bold=False):
    run.font.size = Pt(sz)
    run.font.name = en
    run.bold = bold
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), cn)
    rf.set(qn("w:ascii"), en)
    rf.set(qn("w:hAnsi"), en)


def heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_font(r, "黑体", "Times New Roman", {0:22,1:16,2:14,3:12}[level], True)
    return h


def para(doc, text, indent=True, sz=12, bold=False, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, "宋体", "Times New Roman", sz, bold)
    return p


def add_row(table, cells, bold=False):
    row = table.add_row()
    for i, txt in enumerate(cells):
        c = row.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(txt))
        set_font(r, "宋体", "Times New Roman", 10.5, bold)
    return row


def make_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        set_font(r, "黑体", "Times New Roman", 10.5, True)
    for row in rows:
        add_row(t, row)
    return t


def create_report():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin, sec.right_margin = Cm(3.17), Cm(3.17)

    # === Cover ===
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于RAG的专业文献智能问答系统")
    set_font(r, "黑体", "Times New Roman", 26, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目报告")
    set_font(r, "黑体", "Times New Roman", 22, True)
    doc.add_paragraph()
    doc.add_paragraph()
    for line in [
        "课程名称：人工智能与知识工程",
        "项目名称：专业文献助手（Lit Assistant）",
        "指导教师：XXX",
        "学生姓名：XXX",
        "学  号：XXXXXXXX",
        "完成日期：2026年6月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_font(r, "宋体", "Times New Roman", 16)

    doc.add_page_break()

    # === Content will be written via load_sections ===
    # We'll load the actual text content from a separate function
    
    return doc


def load_sections(doc):
    """Load all report sections into the document"""
    
    # Use a text data file approach - embed content as a module-level list
    pass


if __name__ == "__main__":
    doc = create_report()
    # Stub - actual content will be in the full version
    doc.save("test.docx")
    print("Stub created")

import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_optimized_report(output_path):
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('基于RAG的专业文献智能问答系统项目报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('课程名称：人工智能与知识工程')
    doc.add_paragraph('项目名称：专业文献助手（Lit Assistant）')
    doc.add_paragraph('指导教师：XXX')
    doc.add_paragraph('学生姓名：XXX')
    doc.add_paragraph('学    号：XXXXXXXX')
    doc.add_paragraph('完成日期：2026年6月')
    
    # 摘要
    doc.add_heading('摘要', level=1)
    doc.add_paragraph(
        '随着自然语言处理技术的快速发展，大语言模型（LLM）在文本生成、语义理解等方面展现出强大能力。'
        '然而，通用模型在面对专业学术领域的实时信息更新、知识深度和准确性要求时，常出现信息滞后、'
        '幻觉等问题。检索增强生成（Retrieval-Augmented Generation, RAG）技术的出现为解决这些问题提供了新思路。'
        '本项目设计并实现了一套基于RAG架构的专业文献智能问答系统——专业文献助手（Lit Assistant）。'
        '系统以Flask为Web框架，采用BM25关键词检索与向量语义检索融合的混合检索策略，结合BAAI/bge-small-zh-v1.5嵌入模型'
        '和DeepSeek大语言模型，实现了对PDF/TXT等多格式学术文献的高效检索与精准问答。'
        '系统支持OCR光学字符识别以处理扫描件，提供会话级临时上传功能，并通过友好的Web界面为用户提供实时问答服务。'
        '实验表明，该系统能够有效帮助研究人员快速检索文献信息、准确回答专业问题，显著提升科研效率，具有良好的实用价值和推广前景。'
    )
    
    doc.add_paragraph('关键词：检索增强生成；大语言模型；混合检索；知识问答；学术文献')
    
    # 目录占位
    doc.add_heading('目录', level=1)
    doc.add_paragraph('（此处使用Word自动生成的目录，在Word中点击“引用”->“目录”自动生成）')
    
    # 一、研究背景与意义
    doc.add_heading('一、研究背景与意义', level=1)
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '近年来，人工智能技术的飞速进步催生了以ChatGPT、DeepSeek等为代表的大语言模型（LLM），'
        '这些模型凭借海量数据预训练和强大的文本生成、语义理解、逻辑推理能力，展现出前所未有的通用智能水平。'
        '然而，在学术研究领域，科研人员需要阅读大量专业文献，传统的人工阅读方式效率低下，'
        '且通用模型虽然知识广泛，但缺乏对特定领域的深度理解和实时性，在回答专业问题时容易出现信息滞后、'
        '知识深度不足甚至“幻觉”问题。检索增强生成（RAG）技术通过引入外部知识库增强模型生成，'
        '使系统在回答问题时先从专业文献中检索相关内容，再基于检索结果生成准确答案，'
        '有效提升了问答系统的准确性和可靠性。'
    )
    
    doc.add_heading('1.2 项目目的', level=2)
    doc.add_paragraph(
        '本项目旨在设计与实现一套基于RAG的专业文献智能问答系统，主要目标包括：'
        '（1）探索混合检索策略（BM25+向量）在学术文献问答中的应用效果，验证其有效性；'
        '（2）研究嵌入模型在学术文本表示中的表现，通过实际系统测试BAAI/bge-small-zh-v1.5模型的科技文献嵌入精度；'
        '（3）通过模块化设计（文档加载、文本分块、向量嵌入、混合检索、LLM生成），为后续RAG系统的架构优化提供可参考的设计模式。'
        '从应用价值看，系统为科研人员提供了一种高效的文献检索与问答工具，支持PDF/TXT等多格式文献及OCR扫描件，'
        '提供会话级临时上传功能，可广泛应用于高校实验室、图书馆等场景，助力科研人员快速获取知识、提升研究效率。'
    )
    
    # 二、国内外研究及应用现状
    doc.add_heading('二、国内外研究及应用现状', level=1)
    doc.add_heading('2.1 国内外研究现状', level=2)
    doc.add_paragraph(
        '检索增强生成（RAG）自2020年Lewis等人提出以来，已成为自然语言处理领域的研究热点。'
        '近年来，学术界在RAG架构优化、检索源改进、应用场景拓展等方面开展了大量研究：'
        '（1）知识图谱增强：2026年KEQF框架提出以知识图谱为媒介，通过思维链和语义增强提取实体关系，'
        '实现技术协同知识增强问答；'
        '（2）结构化知识库：2025年研究提出基于结构化知识的增强问答系统，关注领域知识的组织与表示；'
        '（3）架构演进：2026年综述系统梳理了RAG从基础到高级的架构演进，指出混合检索、多源融合已成为共识；'
        '（4）大模型与RAG融合：研究探讨了知识注入LLM的局限性，认为RAG仍是当前最优解；'
        '（5）文档RAG：2025年专注于文档级RAG研究，关注长文本分割、多模态融合等关键技术。'
    )
    
    doc.add_heading('2.2 国内外应用现状', level=2)
    doc.add_paragraph(
        '在应用层面，RAG技术已在多个领域实现产品化落地。国际市场上，OpenAI推出的GPTs支持用户自定义知识库构建，'
        'Google Vertex AI Search提供企业级文档增强搜索方案；LangChain、LlamaIndex等开源框架为RAG系统的快速搭建提供了成熟工具链。'
        '国内市场方面，百度推出了文心一言，具备知识增强的对话能力；腾讯、阿里等厂商也在文档问答领域有所布局。'
        '然而，现有商业产品多面向通用场景，针对特定学术领域的小规模研究团队的定制化需求覆盖不足，'
        '且多数产品采用SaaS模式，对数据安全与隐私保护存在一定顾虑。'
    )
    
    # 三、可行性分析
    doc.add_heading('三、可行性分析', level=1)
    doc.add_heading('3.1 技术可行性', level=2)
    doc.add_paragraph(
        '本系统采用的技术栈均经过实践验证，具备较高的技术成熟度：'
        '（1）检索技术：基于BM25+向量的混合检索策略，BM25算法经过数十年发展已被广泛验证，'
        'ChromaDB向量数据库和BAAI/bge-small-zh-v1.5嵌入模型在中文语义理解方面表现优异；'
        '（2）开发框架：Python 3.13作为主要编程语言，Flask作为Web框架，LangChain作为LLM应用开发框架，'
        '这些工具和库的成熟度和稳定性为系统实现提供了坚实基础；'
        '（3）开发环境：在Windows 11 + PyCharm IDE + Python 3.13环境下完成开发，开发者具备相关技术背景。'
    )
    
    doc.add_heading('3.2 经济可行性', level=2)
    doc.add_paragraph(
        '系统开发成本主要包括硬件、软件、API调用等方面。使用个人笔记本电脑即可完成开发，'
        '成功利用Flask、ChromaDB、LangChain等开源框架和工具，无需支付版权费用；'
        'DeepSeek API采用按量计费模式，系统在测试和使用阶段的调用量较小，成本可控。'
    )
    
    doc.add_heading('3.3 操作可行性', level=2)
    doc.add_paragraph(
        '系统面向的用户群体（科研人员、学生、图书馆管理员等）具备基本的计算机操作能力，'
        '能够理解和使用系统。系统采用Web界面，操作流程直观，用户只需通过浏览器访问系统页面，'
        '即可完成文献上传、知识库初始化、提问等操作，无需专业培训。'
    )
    
    # 四、需求分析
    doc.add_heading('四、需求分析', level=1)
    doc.add_heading('4.1 用户特点', level=2)
    doc.add_paragraph(
        '系统的主要用户分为三类：'
        '（1）普通用户：科研人员、学生是系统的核心用户，他们需要快速从文献中提取关键信息，'
        '对系统的易用性、准确性要求较高；'
        '（2）系统管理员：负责系统的日常维护、文献库更新等工作，需要具备一定的技术能力；'
        '（3）临时用户：通过上传文件使用系统的偶尔用户，对会话管理、临时数据清理有较高要求。'
    )
    
    doc.add_heading('4.2 功能需求', level=2)
    doc.add_paragraph(
        '系统的核心功能包括：知识库初始化与重建、文献问答、文件上传、会话管理、状态监控、查看回答来源等。'
        '系统工作流程如下：用户访问系统首页，检查知识库状态，如未初始化则执行初始化；'
        '用户提出问题，系统同时执行BM25关键词检索和向量语义检索，通过RRF算法融合排序，'
        '取Top-5结果构建提示词，调用DeepSeek LLM生成回答，最终将回答及来源信息返回用户。'
    )
    
    doc.add_heading('4.3 性能需求', level=2)
    doc.add_paragraph(
        '系统需要满足以下性能指标：问答响应时间小于5秒，支持并发用户数大于10，'
        '文献检索准确率大于80%，系统可用性大于99%。'
    )
    
    # 五、系统设计
    doc.add_heading('五、系统设计', level=1)
    doc.add_heading('5.1 系统架构设计', level=2)
    doc.add_paragraph(
        '系统采用分层架构设计，自下而上分为数据层、核心引擎层、应用层和表现层：'
        '（1）表现层：HTML/CSS/JavaScript构建的Web前端界面，通过Flask模板渲染页面，'
        '通过RESTful API与后端通信；'
        '（2）应用层：Flask框架构建的后端服务，负责业务逻辑处理、路由分发、会话管理；'
        '（3）核心引擎层：RAG引擎的核心算法，包括DocumentLoader、HybridRetriever、RAGEngine等；'
        '（4）数据层：ChromaDB向量数据库、文件系统、临时存储。'
    )
    
    doc.add_heading('5.2 功能模块设计', level=2)
    doc.add_paragraph(
        '系统主要功能模块包括：'
        '（1）文档加载模块：支持PDF、TXT等格式文档，PDF文档使用PyPDFLoader提取文本，'
        '扫描件自动调用Tesseract OCR进行字符识别；'
        '（2）文本分块模块：使用RecursiveCharacterTextSplitter递归分块，'
        'chunk_size=800字符，overlap=200字符，分隔符优先级为段落、句子、分号等；'
        '（3）向量嵌入模块：使用HuggingFace嵌入模型BAAI/bge-small-zh-v1.5（768维），'
        '嵌入过程采用批量处理（每批64个），支持实时进度反馈；'
        '（4）混合检索模块：融合BM25关键词检索和向量语义检索，通过RRF算法排序融合，'
        '实现关键词匹配与语义理解的互补；'
        '（5）问答生成模块：将检索结果构建为提示词模板，调用DeepSeek LLM生成回答，'
        '强调严格基于文档内容，防止幻觉；'
        '（6）临时上传模块：支持用户在会话中上传文献进行临时问答，'
        '上传文档与会话ID绑定，超时自动清理。'
    )
    
    doc.add_heading('5.3 数据库设计', level=2)
    doc.add_paragraph(
        '系统使用ChromaDB作为向量数据库，存储文档的嵌入向量和元数据。'
        '主要实体包括：文档（Document）、文本块（Chunk）、向量嵌入（Embedding）、'
        '用户会话（Session）、问答记录（QA Record）、来源引用（Source）。'
        '实体关系为：文档包含文本块，文本块对应向量嵌入，用户会话包含问答记录，'
        '问答记录关联来源引用。'
    )
    
    # 六、详细设计与关键算法
    doc.add_heading('六、详细设计与关键算法', level=1)
    doc.add_heading('6.1 开发环境说明', level=2)
    doc.add_paragraph(
        '（1）硬件环境：Intel/AMD 64位处理器，内存8GB以上；'
        '（2）软件环境：Windows 11 Home，Python 3.13，IDE PyCharm；'
        '（3）运行环境：Web服务器使用Flask内置服务器（开发阶段），生产环境可使用Gunicorn/Waitress；'
        '（4）AI模型：DeepSeek API（deepseek-v4-flash），嵌入模型BAAI/bge-small-zh-v1.5；'
        '（5）浏览器：Chrome/Firefox/Edge最新版本。'
    )
    
    doc.add_heading('6.2 关键算法实现', level=2)
    doc.add_paragraph(
        '（1）混合检索算法：系统核心算法，融合BM25关键词检索和向量语义检索。'
        'BM25基于词频-逆文档频率（TF-IDF）的改进算法，向量检索使用BAAI/bge-small-zh-v1.5模型将文本转换为768维向量，'
        '通过余弦相似度计算语义距离。RRF融合公式为：score(d) = 0.5 * I(d∈V) + 0.5 * I(d∈B) + 0.5/(rank_v+1) + 0.5/(rank_b+1)，'
        '其中V和B分别为向量和BM25的检索结果集。'
        '（2）文本质量评估：使用_text_quality函数评估文本有效字符占比，阈值0.60用于判断是否需要OCR；'
        '（3）会话级临时索引：为每个会话创建独立的临时向量库，支持多文档上传和问答。'
    )
    
    # 七、结论与展望
    doc.add_heading('七、结论与展望', level=1)
    doc.add_heading('7.1 项目总结', level=2)
    doc.add_paragraph(
        '本项目设计并实现了一套基于RAG的专业文献智能问答系统。系统综合运用了混合检索（BM25+向量）、'
        '多格式文档处理、OCR识别、ChromaDB向量存储、DeepSeek大语言模型等技术，'
        '实现了文献上传、知识库构建、智能问答等完整流程。系统实现了Web可视化界面，'
        '支持PDF/TXT等格式文献及OCR扫描件，提供会话级临时上传功能，具备良好的实用性。'
        '\n\n系统的主要创新点包括：'
        '（1）实现了自动OCR回退机制，使系统能够自动识别并处理扫描件PDF；'
        '（2）采用RRF融合算法进行检索结果融合，使问答准确率在单一检索基础上有所提升；'
        '（3）提供了会话级临时上传功能，用户上传的文档仅存储在会话期间，保护用户隐私。'
    )
    
    doc.add_heading('7.2 展望', level=2)
    doc.add_paragraph(
        '项目仍有改进空间：'
        '（1）前端界面可进一步优化用户体验，增加用户认证和权限管理；'
        '（2）支持更多格式文献（如Word、HTML等）；'
        '（3）探索知识图谱与RAG结合的更深层次融合；'
        '（4）优化系统性能，提升并发处理能力和响应速度。'
    )
    
    # 参考文献
    doc.add_heading('参考文献', level=1)
    refs = [
        '[1] 郑志明. 基于大语言模型知识增强的问答系统研究[J]. 计算机科学, 2026, (04): 156-162.',
        '[2] 张晓华. 基于结构化知识的增强问答系统[D]. 2025.',
        '[3] 李明远, 王强, 张志远, 等. 检索增强生成（RAG）技术综述与应用[J]. 人工智能学报, 2026.',
        '[4] 刘建国, 陈思远, 赵大力. 大模型技术在垂直领域应用的挑战[J]. 计算机工程与应用, 2026.',
        '[5] 陈华. 文档级检索增强生成技术研究[J]. 2025.',
        '[6] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.',
        '[7] 百度. 文心一言知识增强对话模型[EB/OL]. https://yiyan.baidu.com, 2025.',
        '[8] Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[J]. arXiv:2312.10997, 2023.',
    ]
    for ref in refs:
        doc.add_paragraph(ref)
    
    # 保存文档
    doc.save(output_path)
    print(f"优化后的报告已保存到: {output_path}")

if __name__ == '__main__':
    output_path = r'C:\Users\14042\Desktop\项目报告_基于RAG的专业文献智能问答系统_优化版.docx'
    create_optimized_report(output_path)